import openai
import os
import shutil
import argparse
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.document_loaders.pdf import PyPDFDirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate
from docx import Document

load_dotenv()
os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")

DATA_PATH = "data"
CHROMA_PATH = "chroma"

PROMPT_TEMPLATE = """
Based on the following context, generate a concise 2-page summary that includes the following sections:

1. **Business Overview**: Include information such as the company's formation/incorporation date, headquarters location, business description, employee count, latest revenues, stock exchange listing and market capitalization, number of offices and locations, and details on their clients/customers.

2. **Business Segment Overview**: Extract the revenue percentage of each component (verticals, products, segments, and sections) as a part of the total revenue. Evaluate the performance of each component by comparing the current year's sales/revenue and market share with the previous year's numbers. Explain the causes of the increase or decrease in the performance of each component.

3. **Geographical Segment Overview**: Breakdown of sales and revenue by geography, specifying the percentage contribution of each region to the total sales. Summarize geographical data, such as workforce, clients, and offices, and outline the company's regional plans for expansion or reduction. Analyze and explain regional sales fluctuations, including a geographical sales breakdown to identify sales trends.

4. **Year-over-Year Analysis**: Summarize year-over-year sales increase or decline and reasons for the change.

5. **Summary of Rationale & Considerations**: Provide an analysis of the risks and mitigating factors.

6. **SWOT Analysis**: Summarize the company's strengths, weaknesses, opportunities, and threats.

7. **Credit Rating Information**: Provide information about credit rating/credit rating changes/changes in the rating outlook.

The 1-page summary should be a condensed version of the 2-pager, including key numbers and important points from the business segment overview and geographical segment overview.

Context:
{context}

---

Generate the 2-page summary based on the above context.
"""

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("query_text", type=str, help="The query text.")
    args = parser.parse_args()
    query_text = args.query_text
    documents = load_documents()
    chunks = split_documents(documents)
    
    embed_model = OpenAIEmbeddings(model="text-embedding-3-large")
    save_to_chroma(chunks, embed_model)
    
    response_text = query_rag(query_text)
    save_to_docx(response_text, "2-page-summary.docx", "1-page-summary.docx")

def load_documents():
    document_loader = PyPDFDirectoryLoader(DATA_PATH)
    return document_loader.load()

def split_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=80,
        length_function=len,
        is_separator_regex=False,
    )
    return text_splitter.split_documents(documents)

def save_to_chroma(chunks, embed_model):
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)

    db = Chroma.from_documents(
        chunks, embedding=embed_model, persist_directory=CHROMA_PATH
    )
    db.persist()
    print(f"Saved {len(chunks)} chunks to {CHROMA_PATH}.")

    chunks_with_ids = calculate_chunk_ids(chunks)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])
    print(f"Number of existing documents in DB: {len(existing_ids)}")

    new_chunks = []
    for chunk in chunks_with_ids:
        if chunk.metadata["id"] not in existing_ids:
            new_chunks.append(chunk)

    if len(new_chunks):
        new_chunk_ids = [chunk.metadata["id"] for chunk in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
        db.persist()
    else:
        print("No new documents to add")

def calculate_chunk_ids(chunks):
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source")
        page = chunk.metadata.get("page")
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk_id = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id
        chunk.metadata["id"] = chunk_id

    return chunks

def clear_database():
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)

def query_rag(query_text: str):
    embedding_function = embed_model
    db = Chroma(persist_directory=CHROMA_PATH, embedding_function=embedding_function)

    results = db.similarity_search_with_score(query_text, k=5)

    context_text = "\n\n---\n\n".join([doc.page_content for doc, _score in results])
    prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
    prompt = prompt_template.format(context=context_text, question=query_text)

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )

    response_text = response['choices'][0]['message']['content']

    sources = [doc.metadata.get("id", None) for doc, _score in results]
    formatted_response = f"Response: {response_text}\nSources: {sources}"
    print(formatted_response)
    return response_text

def save_to_docx(response_text, file_2page, file_1page):
    doc = Document()
    doc.add_heading('2-Page Summary', 0)
    doc.add_paragraph(response_text)
    doc.save(file_2page)

    # Generate 1-page summary
    prompt_template = """
    Based on the following 2-page summary, generate a concise 1-page summary, including key numbers and important points from the business segment overview and geographical segment overview:

    2-Page Summary:
    {context}

    ---

    Generate the 1-page summary based on the above context.
    """
    prompt = prompt_template.format(context=response_text)
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )

    response_text_1page = response['choices'][0]['message']['content']

    doc = Document()
    doc.add_heading('1-Page Summary', 0)
    doc.add_paragraph(response_text_1page)
    doc.save(file_1page)

if __name__ == "__main__":
    main()