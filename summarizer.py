import PyPDF2
import spacy
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import openai
from docx import Document

# Set up OpenAI API key
openai.api_key = 'YOUR_API_KEY'

# Load spaCy model
nlp = spacy.load('en_core_web_sm')

# Load SentenceTransformer model
model = SentenceTransformer('all-MiniLM-L6-v2')

# Extract text from PDF
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
    return text

# Preprocess text
def preprocess_text(text):
    doc = nlp(text)
    return ' '.join([token.lemma_ for token in doc if not token.is_stop and not token.is_punct])

# Extract Business Overview
def extract_business_overview(text):
    doc = nlp(text)
    business_overview = {
        "formation_date": "",
        "headquarters_location": "",
        "business_description": "",
        "employee_count": "",
        "latest_revenues": "",
        "stock_exchange_listing": "",
        "market_capitalization": "",
        "number_of_offices": "",
        "clients_customers": ""
    }

    # Extract relevant details using NER and regex patterns
    for ent in doc.ents:
        if ent.label_ == "GPE":
            business_overview["headquarters_location"] = ent.text
            break

    return business_overview

# Extract Business Segment Overview
def extract_business_segment_overview(text):
    doc = nlp(text)
    segments = {}
    # Define rules or patterns to extract segments
    for ent in doc.ents:
        if ent.label_ == "ORG":
            segments[ent.text] = {"revenue_percentage": "", "performance": {}}
    return segments

# Extract Geographical Breakdown
def extract_geographical_breakdown(text):
    doc = nlp(text)
    geography = {}
    # Extract geographical data
    for ent in doc.ents:
        if ent.label_ == "GPE":
            geography[ent.text] = {"sales_percentage": "", "workforce": "", "clients": "", "offices": ""}
    return geography

# Generate embeddings
def get_embeddings(texts):
    return model.encode(texts)

# Store embeddings in vector database
def store_embeddings(embeddings):
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings)
    return index

# Retrieve relevant chunks
def search_similar_chunks(index, query_embedding, k=5):
    D, I = index.search(np.array([query_embedding]), k)
    return I

# Summarize text with OpenAI API
def summarize_text(text_chunks):
    summaries = []
    for chunk in text_chunks:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "system", "content": "Summarize the following text."},
                      {"role": "user", "content": chunk}]
        )
        summaries.append(response['choices'][0]['message']['content'])
    return summaries

# Generate summary document
def create_summary_doc(summary, output_path):
    doc = Document()
    doc.add_heading('Financial Report Summary', 0)
    for section_title, section_content in summary.items():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(section_content)
    doc.save(output_path)

# Main workflow
def main(pdf_path, output_path):
    text = extract_text_from_pdf(pdf_path)
    preprocessed_text = preprocess_text(text)

    business_overview = extract_business_overview(preprocessed_text)
    business_segment_overview = extract_business_segment_overview(preprocessed_text)
    geographical_breakdown = extract_geographical_breakdown(preprocessed_text)

    text_chunks = [business_overview, business_segment_overview, geographical_breakdown]
    embeddings = get_embeddings(text_chunks)
    index = store_embeddings(embeddings)

    relevant_chunks = []
    for chunk in text_chunks:
        query_embedding = get_embeddings([chunk])[0]
        chunk_indices = search_similar_chunks(index, query_embedding)
        relevant_chunks.append(text_chunks[chunk_indices[0]])

    summaries = summarize_text(relevant_chunks)

    summary = {
        "Business Overview": summaries[0],
        "Business Segment Overview": summaries[1],
        "Geographical Breakdown": summaries[2]
    }
    create_summary_doc(summary, output_path)

if __name__ == "__main__":
    main("path/to/financial_report.pdf", "path/to/summary.docx")